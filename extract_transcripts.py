#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Extract transcript texts from .docx, .txt, and .pdf files"""

import docx
import json
import logging
from pathlib import Path
from typing import Dict, List
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import datetime
from tqdm import tqdm

# Setup logging (will be configured in main() after creating output dir)
logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {'.docx', '.txt', '.pdf', '.srt', '.vtt'}


def extract_text(file_path: Path) -> str:
    """Extract text from supported file types"""
    ext = file_path.suffix.lower()

    if ext == '.docx':
        return _extract_docx(file_path)
    elif ext == '.txt':
        return _extract_txt(file_path)
    elif ext == '.pdf':
        return _extract_pdf(file_path)
    elif ext in ('.srt', '.vtt'):
        return _extract_subtitle(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _extract_docx(file_path: Path) -> str:
    """Extract text from Word document"""
    try:
        doc = docx.Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        table_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                table_text.append(' | '.join(row_text))

        return '\n\n'.join(paragraphs + table_text)
    except Exception as e:
        logger.exception(f"提取失败 {file_path}")
        raise


def _extract_txt(file_path: Path) -> str:
    """Extract text from plain text file"""
    try:
        return file_path.read_text(encoding='utf-8')
    except UnicodeDecodeError:
        return file_path.read_text(encoding='gbk')


def _extract_pdf(file_path: Path) -> str:
    """Extract text from PDF file"""
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("PDF support requires pdfplumber. Run: pip install pdfplumber")

    text_parts = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

    return '\n\n'.join(text_parts)


def _extract_subtitle(file_path: Path) -> str:
    """Extract text from SRT/VTT subtitle file, stripping timestamps"""
    import re

    try:
        content = file_path.read_text(encoding='utf-8')
    except UnicodeDecodeError:
        content = file_path.read_text(encoding='gbk')

    # Remove WEBVTT header
    content = re.sub(r'^WEBVTT.*?\n', '', content)

    # Remove sequence numbers (lines that are just digits)
    content = re.sub(r'^\d+\s*$', '', content, flags=re.MULTILINE)

    # Remove timestamp lines (00:00:00,000 --> 00:00:02,000 or similar)
    content = re.sub(r'\d{2}:\d{2}:\d{2}[.,]\d{3}\s*-->\s*\d{2}:\d{2}:\d{2}[.,]\d{3}.*', '', content)

    # Clean up extra blank lines
    content = re.sub(r'\n{3,}', '\n\n', content)

    return content.strip()


def scan_source_files(source_dir: Path) -> List[Path]:
    """Scan directory for all supported transcript files"""
    files = []
    for ext in SUPPORTED_EXTENSIONS:
        files.extend(source_dir.glob(f'*{ext}'))
    return sorted(files, key=lambda p: p.name)

def main() -> None:
    """Extract transcripts and save to JSON"""
    import sys
    sys.stdout.reconfigure(encoding='utf-8')

    # Use relative paths
    project_root = Path(__file__).parent
    source_dir = project_root / "source_documents"
    output_dir = project_root / "output"

    # Ensure output directory exists
    output_dir.mkdir(exist_ok=True)

    # Setup logging (after output dir exists)
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(levelname)s - %(message)s',
        handlers=[
            logging.FileHandler(output_dir / f'extraction_{datetime.now():%Y%m%d_%H%M%S}.log', encoding='utf-8'),
            logging.StreamHandler()
        ]
    )

    # Define transcript groups - using EXACT filenames from disk
    # Note: Files use full-width Chinese punctuation (：，？！) not ASCII (:,?!)
    groups = {
        "第一章：冷战与大航海时代": [
            "已核对：1、冷战（序言01）麦哲伦的条件.docx",
            "已核对：第2集.docx",
            "已核对：3、冷战，航海时代，澳大利亚全球战略的背后，美国的冷战思维.docx",
            "已核对：4、冷战，大航海，谁的民主？谁的麦哲伦？儒 家的郑和，大明的宝船舰队.docx",
            "已核对：5、冷战大航海时代、郑和不殖民.docx",
            "已核对：6、冷战，大航海时代，美国打到一半的文明土著.docx",
            "已核对：7、冷战，大航海-儒家思想与教会.docx",
            "已核对：8 大航海与冷战.docx"
        ],
        "第二章：黄金与比特币": [
            "金与人民币，历史地位？（相关性的一些思考）20231210直播课.docx",
            "比特币ETF合法化与黄金的宏观联系20240121直播课.docx",
            "20240315黄金与比特币ETF的双龙戏凤格局（上篇）.docx",
            "20240322黄金与比特币ETF下篇.docx",
            "20240329黄金，美元比特币课后答疑.docx",
            "金，2024-2025要回调？ 20240608直播课.docx",
            "答疑总结：金预判，美股分化与央行博弈20240614直播课.docx",
            "20241123货币新时代的金油汇债股数字货币等详解(逐字稿).docx",
            "20260111直播课 2026年：未来36个月 金 油 汇 币 股 总纲 三年后再看 续篇.docx"
        ],
        "第三章：估值模型与资本演变": [
            "20240327直播课：资本千年演变（续）：看懂当下的特斯拉与美股三大估值模型.docx",
            "茅台 大师兄从批到捧，为什么.docx",
            "20241218从阿根廷到日本，从米莱到川马同行，从圈地运动到资本政府1000年.docx",
            "20250612黑盒子估值模型（第二期）：未来三年美联储的困局，华夏消费，黄金，虚拟货币.docx",
            "20250621-2025禁酒令：茅台估值，泡泡玛特估值与资产负债表修复.docx",
            "已核对20250927直播课：稳定币货币承载与未来资产估值.docx"
        ],
        "第四章：中国股市与宏观经济": [
            "后疫情时代的宏观经济.docx",
            "美国房、中国地产的宏观经济启示.docx",
            "书呆子补课（20240915直播课的补课）.docx",
            "大A的明天？答案尽在昨天和今天！（上篇）20240915直播课.docx",
            "2025.11.23大A（中篇）：历史转折点！？（修订版）.docx"
        ],
        "第五章：科技投资与数字经济": [
            "大师兄视频课程20210808.docx",
            "大师兄 关于石油新能源的讨论20210813.docx",
            "北交所与未来科技投资.docx",
            "数字经济2022.docx"
        ],
        "第六章：文化传承与人生哲学": [
            "20231231直播课：开启新的一年！悟而后醒2024元旦祝语.docx",
            "20240112元旦祝语二：苏东坡，弟子规，道德经.docx",
            "从梁思成与林徽因的教育与婚姻，看华夏了不起的传承20250215直播课.docx"
        ],
        "第七章：投资心法与健康": [
            "大师兄聊神归与投资与疾病的关系20210830.docx"
        ]
    }

    logger.info("开始提取讲座文稿...")

    # Prepare all tasks for parallel processing
    tasks = []
    for chapter, files in groups.items():
        for filename in files:
            tasks.append((chapter, filename, source_dir / filename))

    # Extract all transcripts in parallel
    all_data: Dict[str, Dict] = {chapter: {} for chapter in groups.keys()}

    with ThreadPoolExecutor(max_workers=4) as executor:
        # Submit all tasks
        futures = {
            executor.submit(extract_text, filepath): (chapter, filename)
            for chapter, filename, filepath in tasks
        }

        # Process results with progress bar
        for future in tqdm(as_completed(futures), total=len(futures), desc="提取进度"):
            chapter, filename = futures[future]
            try:
                text = future.result()
                all_data[chapter][filename] = {
                    "text": text,
                    "char_count": len(text),
                    "word_estimate": len(text) // 2  # Rough estimate for Chinese
                }
                logger.debug(f"成功提取: {filename}")
            except FileNotFoundError:
                logger.warning(f"文件不存在: {filename}")
                all_data[chapter][filename] = {"error": "File not found"}
            except Exception as e:
                logger.warning(f"提取失败 {filename}: {e}")
                all_data[chapter][filename] = {"error": str(e)}

    # Save to JSON with readable formatting
    output_file = output_dir / "transcripts_extracted.json"
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(all_data, f, ensure_ascii=False, indent=2)

    logger.info(f"\n所有文本已提取到: {output_file}")

    # Print summary
    print("\n=== 提取摘要 ===")
    total_all_chars = 0
    total_all_files = 0
    for chapter, chapter_data in all_data.items():
        total_chars = sum(d.get('char_count', 0) for d in chapter_data.values())
        file_count = len(chapter_data)
        total_all_chars += total_chars
        total_all_files += file_count
        print(f"{chapter}: {file_count} 个文件, 共 {total_chars:,} 字")

    print(f"\n总计: {total_all_files} 个文件, 共 {total_all_chars:,} 字")

if __name__ == "__main__":
    main()
